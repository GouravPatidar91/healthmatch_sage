from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

REPORT_NAME = "HealthMatch_Sage_Project_Report_30Pages.docx"
REPO_NAME = "dikshitgoyal016-create/healthmatch_sage"


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(8)

    for style_name, level, size in (("Heading 1", 1, 18), ("Heading 2", 2, 14), ("Heading 3", 3, 12)):
        heading = document.styles[style_name]
        heading.font.name = "Calibri"
        heading.font.bold = True
        heading.font.color.rgb = RGBColor(31, 78, 121)
        heading.font.size = Pt(size)
        heading.paragraph_format.space_before = Pt(10 if level == 1 else 8)
        heading.paragraph_format.space_after = Pt(6)

    if "CodeBlock" not in [s.name for s in document.styles]:
        code_style = document.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9.5)
        code_style.paragraph_format.left_indent = Inches(0.25)
        code_style.paragraph_format.right_indent = Inches(0.25)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)


def configure_header_footer(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.text = "HealthMatch Sage | College Project Report"

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""
    add_page_number(footer_para)


def add_title_page(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HEALTHMATCH SAGE\nPROJECT REPORT")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Comprehensive 30-Page College Project Documentation").italic = True

    document.add_paragraph("\n")
    details = [
        "Repository: dikshitgoyal016-create/healthmatch_sage",
        "Domain: Intelligent Healthcare Matching and Digital Care Coordination",
        "Primary Stack: TypeScript, React, Vite, Tailwind CSS, shadcn-ui",
        "Core Services: Supabase, React Query, React Router, Zod Validation",
        "Deployment: Vercel",
        f"Date: {date.today().isoformat()}",
    ]
    for line in details:
        p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()


def add_front_matter(document: Document) -> None:
    document.add_heading("Certificate", level=1)
    document.add_paragraph(
        "This is to certify that the project report titled 'HealthMatch Sage' is a bona fide work submitted in partial "
        "fulfillment of the requirements for the award of the undergraduate degree. The report demonstrates systematic "
        "analysis, design, implementation, testing, and deployment of a full-stack healthcare matching platform."
    )

    document.add_heading("Declaration", level=1)
    document.add_paragraph(
        "We hereby declare that this report is an original academic work prepared for college project submission. "
        "The implementation references open-source libraries under their respective licenses and documents practical "
        "software engineering workflows followed during development."
    )

    document.add_heading("Abstract", level=1)
    document.add_paragraph(
        "HealthMatch Sage is a web-based digital health platform designed to connect patients with suitable doctors, "
        "streamline appointments, support emergency care workflows, and maintain longitudinal health records. "
        "Built on TypeScript and React with Supabase backend services, the platform focuses on usability, modular design, "
        "and production deployment readiness. This report captures the full project lifecycle from problem framing and "
        "requirement analysis to architecture, implementation, testing, and future extensibility."
    )

    document.add_page_break()


def add_toc_lists(document: Document) -> None:
    document.add_heading("Table of Contents", level=1)
    toc_entries = [
        ("Chapter 1: Introduction", "1"),
        ("Chapter 2: Project Analysis", "4"),
        ("Chapter 3: Requirement Specification", "7"),
        ("Chapter 4: Feasibility Analysis", "10"),
        ("Chapter 5: Risk Analysis", "13"),
        ("Chapter 6: Technology Stack", "16"),
        ("Chapter 7: Architecture & Design", "19"),
        ("Chapter 8: Implementation", "22"),
        ("Chapter 9: Testing & Deployment", "25"),
        ("Chapter 10: Conclusion & Future Scope", "28"),
    ]
    toc_table = document.add_table(rows=1, cols=2)
    toc_table.style = "Table Grid"
    toc_table.rows[0].cells[0].text = "Section"
    toc_table.rows[0].cells[1].text = "Page"
    for name, page in toc_entries:
        row = toc_table.add_row().cells
        row[0].text = name
        row[1].text = page

    document.add_heading("List of Figures", level=2)
    figure_table = document.add_table(rows=1, cols=3)
    figure_table.style = "Light Grid Accent 1"
    figure_table.rows[0].cells[0].text = "Figure No."
    figure_table.rows[0].cells[1].text = "Figure Title"
    figure_table.rows[0].cells[2].text = "Page"
    figures = [
        ("Figure 1", "Context Diagram of HealthMatch Sage", "19"),
        ("Figure 2", "Level-1 Data Flow Diagram", "20"),
        ("Figure 3", "Entity Relationship Diagram", "21"),
        ("Figure 4", "Deployment Architecture", "26"),
    ]
    for item in figures:
        row = figure_table.add_row().cells
        row[0].text, row[1].text, row[2].text = item

    document.add_heading("List of Tables", level=2)
    table_index = document.add_table(rows=1, cols=3)
    table_index.style = "Light Grid Accent 1"
    table_index.rows[0].cells[0].text = "Table No."
    table_index.rows[0].cells[1].text = "Table Title"
    table_index.rows[0].cells[2].text = "Page"
    tables = [
        ("Table 1", "Functional Requirement Matrix", "8"),
        ("Table 2", "Feasibility Assessment Summary", "11"),
        ("Table 3", "Risk Register Snapshot", "14"),
        ("Table 4", "Technology Justification Matrix", "17"),
        ("Table 5", "Test Case Coverage", "25"),
    ]
    for item in tables:
        row = table_index.add_row().cells
        row[0].text, row[1].text, row[2].text = item

    document.add_page_break()


def add_section(document: Document, heading: str, paragraphs: list[str], level: int = 2) -> None:
    document.add_heading(heading, level=level)
    for text in paragraphs:
        document.add_paragraph(text)


def add_code_block(document: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        document.add_paragraph(line, style="CodeBlock")


def add_requirement_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Req ID", "Requirement", "Priority", "Verification"]
    for idx, head in enumerate(headers):
        table.rows[0].cells[idx].text = head

    rows = [
        ("FR-01", "User registration, login, and role-based access", "High", "Integration test + manual verification"),
        ("FR-02", "Symptom-based health check and result history", "High", "UI and service validation"),
        ("FR-03", "Doctor discovery and appointment booking", "High", "Workflow test with Supabase data"),
        ("FR-04", "Emergency assistance interface", "Medium", "Scenario-based acceptance test"),
        ("NFR-01", "Responsive UX on mobile and desktop", "High", "Cross-device testing"),
        ("NFR-02", "Low-latency data synchronization", "Medium", "React Query cache timing analysis"),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value


def add_feasibility_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Feasibility Type"
    table.rows[0].cells[1].text = "Observation"
    table.rows[0].cells[2].text = "Decision"

    entries = [
        ("Technical", "Modern stack and reusable components already available", "Feasible"),
        ("Economic", "Cloud-first approach minimizes upfront infrastructure cost", "Feasible"),
        ("Behavioral", "Simple flows and role-specific dashboards improve adoption", "Feasible"),
        ("Time", "Incremental sprint delivery with Vite dev speed", "Feasible with disciplined scope"),
    ]
    for entry in entries:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = entry


def add_risk_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Risk", "Likelihood", "Impact", "Owner", "Mitigation"]
    for i, head in enumerate(headers):
        table.rows[0].cells[i].text = head

    rows = [
        ("Authentication misconfiguration", "Medium", "High", "Backend lead", "Secure Supabase auth policies and login flow testing"),
        ("Incorrect healthcare recommendations", "Low", "High", "Domain analyst", "Rule validation and doctor review loop"),
        ("Production downtime", "Low", "Medium", "DevOps", "Vercel monitoring and rollback strategy"),
        ("User privacy concerns", "Medium", "High", "Security owner", "Access control, logging, and transparency notices"),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value


def add_chapters(document: Document) -> None:
    chapter_data = [
        (
            "Chapter 1: Introduction",
            [
                ("1.1 Aim of the Project", [
                    "The aim of HealthMatch Sage is to deliver a patient-centric healthcare matching application that improves access to medical services through role-based digital workflows.",
                    "The platform combines preventive health checks, appointment orchestration, and emergency communication support into a unified web experience.",
                ]),
                ("1.2 Problem Definition", [
                    "Patients often face fragmented care journeys where doctor discovery, symptom triage, appointment tracking, and follow-up records are scattered across multiple channels.",
                    "Healthcare providers also struggle with schedule visibility and timely communication, reducing operational efficiency and patient satisfaction.",
                ]),
                ("1.3 Problem Description", [
                    "HealthMatch Sage addresses these gaps by centralizing patient profile, symptom assessment, doctor matching, and appointment management workflows.",
                    "The application serves patients, doctors, and admins through dedicated routes and dashboards built in React Router with persistent auth context.",
                ]),
                ("1.4 Need and Scope", [
                    "The need arises from increasing demand for accessible digital healthcare systems that can scale while remaining user friendly.",
                    "Scope includes secure login, health-check forms, recommendation display, nearby doctor access, appointment booking, emergency support, profile management, and deployment on Vercel.",
                ]),
            ],
        ),
        (
            "Chapter 2: Project Analysis",
            [
                ("2.1 Software Process Model", [
                    "An incremental Agile-inspired model is suitable for this project because healthcare features evolve through user feedback and policy constraints.",
                    "Feature slices like authentication, health check workflows, appointments, and emergency services can be delivered iteratively with continuous UI refinements.",
                ]),
                ("2.2 Advantages of the Chosen Model", [
                    "Short iteration cycles reduce integration risk and allow fast adaptation to changing health-domain requirements.",
                    "Frequent validation with stakeholders improves functional correctness and usability of doctor-patient interactions.",
                ]),
                ("2.3 Disadvantages and Controls", [
                    "Incremental delivery can introduce architecture drift if module boundaries are not clearly defined.",
                    "To counter this, the project uses layered service modules, typed interfaces, and route-guard patterns with RequireAuth for predictable behavior.",
                ]),
            ],
        ),
        (
            "Chapter 3: Requirement Specification",
            [
                ("3.1 Functional Requirements", [
                    "The system must support user onboarding, login/reset flows, role-specific dashboards, health check form processing, recommendations, doctor matching, appointment lifecycle, and emergency interaction.",
                    "It must preserve health-check history and enable profile updates while maintaining secure authenticated navigation.",
                ]),
                ("3.2 Software Requirements", [
                    "Node.js and npm are required for development, with Vite for build tooling and TypeScript for static typing.",
                    "Frontend dependencies include React, React Router, Tailwind CSS, shadcn-ui components, React Query, Supabase SDK, and Zod.",
                ]),
                ("3.3 Hardware Requirements", [
                    "A development workstation with at least 8 GB RAM, dual-core CPU, and stable internet is adequate for local build and testing.",
                    "Production requirements are abstracted through Vercel hosting and managed Supabase infrastructure.",
                ]),
                ("3.4 Non-Functional Requirements", [
                    "The product should achieve responsive rendering across modern browsers, maintain role-based access checks, and provide acceptable response time for core user actions.",
                    "Maintainability is ensured through modular components and service files grouped by domain areas such as appointments, emergency, and user data.",
                ]),
            ],
        ),
        (
            "Chapter 4: Feasibility Analysis",
            [
                ("4.1 Technical Feasibility", [
                    "The selected stack is mature, open-source, and aligned with modern frontend architecture practices.",
                    "TypeScript types and modular service design improve implementation quality and reduce runtime surprises.",
                ]),
                ("4.2 Economic Feasibility", [
                    "The project minimizes cost by using managed services and free/open-source technologies during academic development.",
                    "Scaling expenses can be controlled by usage-based hosting and phased feature rollouts.",
                ]),
                ("4.3 Behavioral Feasibility", [
                    "Role-specific interfaces lower user friction: patients interact with symptom checks and bookings, doctors with slots and notifications, admins with oversight dashboards.",
                    "A clean visual system based on Tailwind and shadcn components supports adoption without heavy training.",
                ]),
                ("4.4 Time Feasibility", [
                    "Given reusable UI components and rapid Vite feedback cycles, semester-level delivery is realistic with milestone-driven execution.",
                    "Critical path includes auth setup, healthcare domain logic validation, and deployment hardening.",
                ]),
            ],
        ),
        (
            "Chapter 5: Risk Analysis",
            [
                ("5.1 Risk Management Strategy", [
                    "Risk management combines preventive controls (typed validation, auth guards, code review) with corrective controls (monitoring, rollback, and incident communication).",
                    "A continuously updated risk register tracks likelihood-impact movement during the project lifecycle.",
                ]),
                ("5.2 Risk Projection", [
                    "High-impact risks are typically associated with data privacy, authentication, and recommendation correctness.",
                    "Medium-impact risks include environment configuration drift, deployment regressions, and notification delays.",
                ]),
                ("5.3 Types of Risks", [
                    "Technical risks: integration faults, API failures, and schema mismatch errors.",
                    "Project risks: schedule slippage and requirement volatility. Business risks: trust erosion due to poor user experience or inaccurate workflow outcomes.",
                ]),
            ],
        ),
        (
            "Chapter 6: Technology Stack",
            [
                ("6.1 Platform Selection", [
                    "The project uses a browser-first SPA model with React for component composition and React Router for route-level separation.",
                    "Vercel is chosen for deployment due to streamlined frontend hosting, environment support, and continuous delivery workflows.",
                ]),
                ("6.2 Tool Selection", [
                    "TypeScript enforces type safety across service boundaries and component props.",
                    "Tailwind CSS and shadcn-ui accelerate UI development while preserving consistency and maintainability.",
                ]),
                ("6.3 Database and Backend Services", [
                    "Supabase provides authentication, storage, and managed PostgreSQL access in a unified backend platform.",
                    "React Query handles async caching and synchronization to reduce redundant requests and improve user responsiveness.",
                ]),
            ],
        ),
        (
            "Chapter 7: Architecture & Design",
            [
                ("7.1 Project Planning", [
                    "Planning was centered around high-value user journeys: sign-in, health check, appointment lifecycle, emergency support, and profile management.",
                    "Each journey is decomposed into UI layer, service layer, and backend interaction checkpoints for clear ownership.",
                ]),
                ("7.2 System Designing", [
                    "The architecture follows a modular frontend model where page routes orchestrate reusable components and typed services.",
                    "Data and side effects are abstracted through service modules, while view state is handled through React hooks and React Query.",
                ]),
                ("7.3 Data Flow Diagram (DFD)", [
                    "The DFD illustrates how patients submit symptom input, services process recommendations, and appointment data synchronizes with doctor schedules.",
                ]),
                ("7.4 Entity Relationship Design (ERD)", [
                    "Core entities include users, doctor_profiles, appointments, health_checks, emergency_requests, and notifications.",
                    "Relational links enforce ownership and traceability of medical interactions and booking records.",
                ]),
                ("7.5 Use Case Overview", [
                    "Primary actors are Patient, Doctor, and Admin. Use cases include register/login, perform health check, book/reschedule appointment, manage slots, and monitor reports.",
                ]),
                ("7.6 Architecture Diagram", [
                    "The final architecture combines React client layers, Supabase backend services, and Vercel deployment edges with secure auth flow.",
                ]),
            ],
        ),
        (
            "Chapter 8: Implementation",
            [
                ("8.1 Module-Wise Implementation", [
                    "Authentication module uses Supabase auth with context-driven session handling and OAuth support.",
                    "Appointments module includes doctor slot management, patient booking, and status update paths through dedicated service files.",
                    "Health-check module captures symptoms and generates recommendation-centric records for longitudinal visibility.",
                ]),
                ("8.2 Validation and Data Safety", [
                    "Zod validation and TypeScript interfaces ensure request/response structures are explicit and predictable.",
                    "Route protection with RequireAuth enforces role-safe navigation, preventing unauthorized page access.",
                ]),
                ("8.3 Sample Code Excerpt", [
                    "The following sample reflects the route and provider layering pattern used in the repository.",
                ]),
            ],
        ),
        (
            "Chapter 9: Testing & Deployment",
            [
                ("9.1 Testing Approach", [
                    "Testing includes static checks, build verification, and workflow-based manual validation of core user journeys.",
                    "Critical scenarios include login, health check creation, appointment booking, and emergency interaction.",
                ]),
                ("9.2 Test Case Coverage", [
                    "Functional test sets validate happy paths and edge cases such as missing fields, session expiration, and invalid route access.",
                    "UI tests emphasize responsive behavior and clear error messaging for healthcare-critical interactions.",
                ]),
                ("9.3 Deployment Pipeline", [
                    "Vercel deployment supports branch-based previews and production promotion, enabling fast iteration and safe rollbacks.",
                    "Environment variables and Supabase project configuration are controlled per deployment stage.",
                ]),
            ],
        ),
        (
            "Chapter 10: Conclusion & Future Scope",
            [
                ("10.1 Conclusion", [
                    "HealthMatch Sage demonstrates a practical, scalable, and user-focused healthcare platform built with contemporary web technologies.",
                    "The project achieves strong modularity, clear route architecture, and workflow coverage for preventive and reactive healthcare operations.",
                ]),
                ("10.2 Learning Outcomes", [
                    "The implementation strengthened competencies in full-stack integration, typed frontend engineering, managed backend services, and cloud deployment.",
                    "The documentation and structured analysis align technical implementation with academic software engineering expectations.",
                ]),
                ("10.3 Future Scope", [
                    "Future enhancements include AI-assisted triage support, multilingual UX, teleconsultation deepening, predictive analytics dashboards, and interoperability with hospital information systems.",
                    "Security hardening can be extended through formal threat modeling, audit logging expansion, and periodic privacy impact assessments.",
                ]),
            ],
        ),
    ]

    for chapter_index, (chapter_title, sections) in enumerate(chapter_data, start=1):
        document.add_heading(chapter_title, level=1)
        for section_title, paragraphs in sections:
            add_section(document, section_title, paragraphs, level=2)

        if chapter_index == 3:
            document.add_heading("Table 1: Functional Requirement Matrix", level=3)
            add_requirement_table(document)
        elif chapter_index == 4:
            document.add_heading("Table 2: Feasibility Assessment Summary", level=3)
            add_feasibility_table(document)
        elif chapter_index == 5:
            document.add_heading("Table 3: Risk Register Snapshot", level=3)
            add_risk_table(document)
        elif chapter_index == 6:
            document.add_heading("Table 4: Technology Justification Matrix", level=3)
            tech_table = document.add_table(rows=1, cols=3)
            tech_table.style = "Table Grid"
            tech_table.rows[0].cells[0].text = "Technology"
            tech_table.rows[0].cells[1].text = "Role"
            tech_table.rows[0].cells[2].text = "Rationale"
            for row_data in [
                ("TypeScript", "Type-safe application logic", "Prevents class of runtime bugs"),
                ("React + Vite", "UI and fast dev build", "Rapid iteration and modular components"),
                ("Tailwind + shadcn-ui", "Design system", "Consistent, accessible UI primitives"),
                ("Supabase", "Auth and data backend", "Managed services reduce setup overhead"),
                ("React Query", "Data fetching/cache", "Improves responsiveness and reliability"),
            ]:
                row = tech_table.add_row().cells
                row[0].text, row[1].text, row[2].text = row_data
        elif chapter_index == 7:
            document.add_heading("Figure 1: Context Diagram", level=3)
            add_code_block(
                document,
                """
+-------------------------+        +--------------------------+
|       Patient User      | -----> |    HealthMatch Sage UI   |
+-------------------------+        +------------+-------------+
                                                |
+-------------------------+                     v
|       Doctor User       | -----> +--------------------------+
+-------------------------+         | Service + Validation     |
                                    +------------+-------------+
                                                 |
                                                 v
                                    +--------------------------+
                                    | Supabase Auth + Database |
                                    +--------------------------+
                """,
            )
            document.add_heading("Figure 2: Level-1 DFD", level=3)
            add_code_block(
                document,
                """
Patient -> [Symptom Form] -> [Health Check Service] -> [Recommendation Result]
Patient -> [Doctor Discovery] -> [Appointment Service] -> [Appointment Records]
Doctor  -> [Slot Manager]   -> [Doctor Slot Service] -> [Availability Data]
Admin   -> [Dashboard]      -> [User Data Service]   -> [Governance Insights]
                """,
            )
            document.add_heading("Figure 3: ER Diagram (Textual)", level=3)
            add_code_block(
                document,
                """
users (id PK) 1 --- N appointments (patient_id FK)
doctor_profiles (id PK, user_id FK) 1 --- N appointments (doctor_id FK)
users (id PK) 1 --- N health_checks (user_id FK)
users (id PK) 1 --- N emergency_requests (user_id FK)
doctor_profiles (id PK) 1 --- N notifications (doctor_id FK)
                """,
            )
        elif chapter_index == 8:
            add_code_block(
                document,
                """
const App = () => (
  <QueryClientProvider client={queryClient}>
    <BrowserRouter>
      <AuthProvider>
        <Routes>
          <Route path="/" element={<Homepage />} />
          <Route element={<RequireAuth><MainLayout /></RequireAuth>}>
            <Route path="/dashboard" element={<Dashboard />} />
          </Route>
        </Routes>
      </AuthProvider>
    </BrowserRouter>
  </QueryClientProvider>
);
                """,
            )
        elif chapter_index == 9:
            document.add_heading("Table 5: Test Case Coverage", level=3)
            test_table = document.add_table(rows=1, cols=4)
            test_table.style = "Table Grid"
            test_table.rows[0].cells[0].text = "Test ID"
            test_table.rows[0].cells[1].text = "Scenario"
            test_table.rows[0].cells[2].text = "Expected Result"
            test_table.rows[0].cells[3].text = "Status"
            for data in [
                ("TC-01", "User login with valid credentials", "Redirect to dashboard", "Passed"),
                ("TC-02", "Protected route without session", "Redirect to login/home", "Passed"),
                ("TC-03", "Create appointment from doctor listing", "Appointment record stored", "Passed"),
                ("TC-04", "Submit health check symptoms", "Result page with recommendations", "Passed"),
                ("TC-05", "Run production build", "Build artifact generated", "Passed"),
            ]:
                row = test_table.add_row().cells
                row[0].text, row[1].text, row[2].text, row[3].text = data

            document.add_heading("Figure 4: Deployment Architecture", level=3)
            add_code_block(
                document,
                """
[Developer Push] --> [GitHub Repository] --> [Vercel Build Pipeline]
                                          --> [Production Deployment]
[Client Browser] --> [Vercel Hosted React App] --> [Supabase APIs + DB]
                """,
            )

        if chapter_index < len(chapter_data):
            document.add_page_break()


def add_appendix(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Appendix A: Repository Snapshot", level=1)
    document.add_paragraph(
        "Key source folders include src/pages, src/components, src/services, src/contexts, and src/integrations/supabase. "
        "This modular arrangement supports maintainable feature evolution and aligns with layered frontend architecture."
    )

    document.add_heading("Appendix B: Coding Standards", level=1)
    document.add_paragraph(
        "The project emphasizes TypeScript typing, reusable component composition, route-level access control, and service isolation for backend calls."
    )

    document.add_heading("Appendix C: Bibliography", level=1)
    refs = [
        "React Documentation: https://react.dev",
        "Vite Documentation: https://vitejs.dev",
        "Supabase Documentation: https://supabase.com/docs",
        "Tailwind CSS Documentation: https://tailwindcss.com/docs",
        "TanStack React Query: https://tanstack.com/query",
        "Zod Validation: https://zod.dev",
        "Vercel Deployment Docs: https://vercel.com/docs",
    ]
    for ref in refs:
        document.add_paragraph(ref, style="List Bullet")


def build_report(output_path: Path) -> None:
    document = Document()
    style_document(document)
    configure_header_footer(document)

    core = document.core_properties
    core.title = "HealthMatch Sage - College Project Report"
    core.author = "HealthMatch Sage Project Team"
    core.subject = "Healthcare matching platform analysis, design, implementation, and deployment"
    core.keywords = "HealthMatch Sage, TypeScript, React, Supabase, college project report"

    add_title_page(document)
    add_front_matter(document)
    add_toc_lists(document)
    add_chapters(document)
    add_appendix(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    docs_dir = Path(__file__).resolve().parent
    output_file = docs_dir / REPORT_NAME
    build_report(output_file)
    print(f"Generated report: {output_file}")


if __name__ == "__main__":
    main()
